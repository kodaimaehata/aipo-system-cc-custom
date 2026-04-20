from __future__ import annotations

from pathlib import Path
import zipfile

import pytest

from docx import Document

from office_automation import word_ops


def _create_sample_docx(path: Path) -> Path:
    document = Document()
    document.core_properties.title = "Quarterly update"
    document.core_properties.author = "Test Runner"
    document.add_paragraph("Intro paragraph")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "Analyst"

    document.add_paragraph("Closing paragraph")
    document.save(path)
    return path


def test_read_returns_ordered_body_paragraph_and_table_data(tmp_path: Path) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")

    result = word_ops.read(source)

    assert result["format"] == "docx"
    assert result["file_path"] == str(source)
    assert result["metadata"]["title"] == "Quarterly update"
    assert result["metadata"]["author"] == "Test Runner"
    assert [item["type"] for item in result["body"]] == ["paragraph", "table", "paragraph"]
    assert result["paragraphs"] == [
        {
            "paragraph_index": 0,
            "body_index": 0,
            "text": "Intro paragraph",
            "style": "Normal",
        },
        {
            "paragraph_index": 1,
            "body_index": 2,
            "text": "Closing paragraph",
            "style": "Normal",
        },
    ]
    assert result["tables"][0]["table_index"] == 0
    assert result["tables"][0]["body_index"] == 1
    assert result["tables"][0]["row_count"] == 2
    assert result["tables"][0]["column_count"] == 2
    assert result["tables"][0]["rows"][1]["cells"] == [
        {"row_index": 1, "column_index": 0, "text": "Alice"},
        {"row_index": 1, "column_index": 1, "text": "Analyst"},
    ]
    assert result["warnings"] == []



def test_read_surfaces_text_box_warnings_when_markup_is_present(tmp_path: Path) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")

    with zipfile.ZipFile(source, mode="a") as archive:
        archive.writestr(
            "word/runtime-warning.xml",
            (
                '<w:root xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:txbxContent />"
                "</w:root>"
            ),
        )

    result = word_ops.read(source)

    assert result["warnings"] == [
        "This document contains Word text box content. SG2 Word V1 only supports body "
        "paragraph and table edits."
    ]



def test_edit_honors_output_path_and_applies_supported_operations(tmp_path: Path) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")
    output_path = tmp_path / "edited" / "result.docx"
    ignored_output_dir = tmp_path / "ignored-output-dir"

    instructions = {
        "operations": [
            {
                "operation": "replace_paragraph_text",
                "paragraph_index": 0,
                "text": "Updated intro paragraph",
            },
            {
                "operation": "insert_paragraph_after",
                "paragraph_index": 0,
                "text": "Inserted paragraph",
            },
            {
                "operation": "delete_paragraph",
                "paragraph_index": 2,
            },
            {
                "operation": "replace_table_cell",
                "table_index": 0,
                "row_index": 1,
                "column_index": 1,
                "text": "Lead Analyst",
            },
            {
                "operation": "append_table_row",
                "table_index": 0,
                "values": ["Bob", "Manager"],
            },
            {
                "operation": "update_table_row",
                "table_index": 0,
                "row_index": 1,
                "values": ["Alicia", "Principal Analyst"],
            },
        ],
        "output_path": str(output_path),
        "output_dir": str(ignored_output_dir),
        "copy_before_edit": True,
        "options": {},
    }

    with pytest.warns(RuntimeWarning, match="may normalize run-level formatting"):
        saved_path = word_ops.edit(source, instructions)

    assert saved_path == output_path
    assert output_path.exists()
    assert not ignored_output_dir.exists()

    updated = word_ops.read(output_path)
    assert [paragraph["text"] for paragraph in updated["paragraphs"]] == [
        "Updated intro paragraph",
        "Inserted paragraph",
    ]
    assert updated["tables"][0]["rows"][1]["cells"] == [
        {"row_index": 1, "column_index": 0, "text": "Alicia"},
        {"row_index": 1, "column_index": 1, "text": "Principal Analyst"},
    ]
    assert updated["tables"][0]["rows"][2]["cells"] == [
        {"row_index": 2, "column_index": 0, "text": "Bob"},
        {"row_index": 2, "column_index": 1, "text": "Manager"},
    ]
    assert word_ops.read(source)["paragraphs"][0]["text"] == "Intro paragraph"



def test_edit_rejects_missing_output_path_runtime_fallback(tmp_path: Path) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")
    output_dir = tmp_path / "copies"

    instructions = {
        "operations": [
            {
                "operation": "replace_table_cell",
                "table_index": 0,
                "row_index": 1,
                "column_index": 0,
                "text": "Copied Alice",
            }
        ],
        "output_dir": str(output_dir),
        "copy_before_edit": True,
        "options": {},
    }

    with pytest.raises(
        ValueError,
        match=r"Word edit instructions must include 'output_path'; runtime save-path fallback is not supported\.",
    ):
        word_ops.edit(source, instructions)



def test_edit_rejects_same_output_path_when_copy_before_edit_is_enabled(tmp_path: Path) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")

    with pytest.raises(ValueError, match="cannot preserve the original"):
        word_ops.edit(
            source,
            {
                "operations": [
                    {
                        "operation": "replace_paragraph_text",
                        "paragraph_index": 0,
                        "text": "Should not save",
                    }
                ],
                "output_path": str(source),
                "copy_before_edit": True,
                "options": {},
            },
        )

    assert word_ops.read(source)["paragraphs"][0]["text"] == "Intro paragraph"



def test_edit_rejects_existing_output_path_when_copy_before_edit_is_disabled(tmp_path: Path) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")
    output_path = _create_sample_docx(tmp_path / "edited.docx")

    with pytest.raises(
        FileExistsError,
        match=r"Word output path '.*edited\.docx' already exists; the runtime will not silently overwrite it\.",
    ):
        word_ops.edit(
            source,
            {
                "operations": [
                    {
                        "operation": "replace_paragraph_text",
                        "paragraph_index": 0,
                        "text": "Should not overwrite",
                    }
                ],
                "output_path": str(output_path),
                "copy_before_edit": False,
                "options": {},
            },
        )

    assert word_ops.read(output_path)["paragraphs"][0]["text"] == "Intro paragraph"



def test_edit_rejects_existing_copy_before_edit_output_path_for_same_name_in_new_directory(
    tmp_path: Path,
) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")
    copied_output = tmp_path / "copies" / source.name
    copied_output.parent.mkdir()
    _create_sample_docx(copied_output)

    with pytest.raises(
        FileExistsError,
        match=r"Word output path '.*sample\.docx' already exists; the runtime will not silently overwrite it\.",
    ):
        word_ops.edit(
            source,
            {
                "operations": [
                    {
                        "operation": "replace_table_cell",
                        "table_index": 0,
                        "row_index": 1,
                        "column_index": 0,
                        "text": "Should not overwrite",
                    }
                ],
                "output_path": str(copied_output),
                "copy_before_edit": True,
                "options": {},
            },
        )

    assert word_ops.read(copied_output)["tables"][0]["rows"][1]["cells"][0]["text"] == "Alice"



def test_edit_uses_copy_helper_path_when_output_path_keeps_source_name_in_new_directory(tmp_path: Path) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")
    copied_output = tmp_path / "copies" / source.name

    instructions = {
        "operations": [
            {
                "operation": "replace_table_cell",
                "table_index": 0,
                "row_index": 1,
                "column_index": 0,
                "text": "Copied Alice",
            }
        ],
        "output_path": str(copied_output),
        "copy_before_edit": True,
        "options": {},
    }

    with pytest.warns(RuntimeWarning, match="may normalize run-level formatting"):
        saved_path = word_ops.edit(source, instructions)

    assert saved_path == copied_output
    assert saved_path.exists()
    assert word_ops.read(saved_path)["tables"][0]["rows"][1]["cells"][0]["text"] == "Copied Alice"
    assert word_ops.read(source)["tables"][0]["rows"][1]["cells"][0]["text"] == "Alice"



def test_edit_rejects_track_change_cleanup_requests(tmp_path: Path) -> None:
    source = _create_sample_docx(tmp_path / "sample.docx")
    output_path = tmp_path / "edited.docx"

    instructions = {
        "operations": [{"operation": "accept_track_changes"}],
        "output_path": str(output_path),
        "copy_before_edit": True,
        "options": {},
    }

    with pytest.raises(
        NotImplementedError,
        match="Track changes acceptance/rejection is outside Word V1 runtime scope",
    ):
        word_ops.edit(source, instructions)
