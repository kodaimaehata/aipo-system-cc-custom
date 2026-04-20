from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document
from PIL import Image

from office_automation.common.images import extract_images, replace_image


_IMAGE_SIZE = (40, 40)


def _create_color_image(path: Path, color: tuple[int, int, int]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", _IMAGE_SIZE, color)
    image.save(path, format="PNG")
    return path


def _image_color(path: Path) -> tuple[int, int, int]:
    with Image.open(path) as image:
        return image.convert("RGB").getpixel((0, 0))


def _create_docx_with_images(path: Path, image_paths: list[Path]) -> Path:
    document = Document()
    document.add_paragraph("Image sample")
    for image_path in image_paths:
        document.add_picture(str(image_path))
    document.save(path)
    return path


def _create_pdf_with_images(path: Path, image_paths: list[Path]) -> Path:
    document = fitz.open()
    first_page = document.new_page()
    first_page.insert_image(fitz.Rect(36, 36, 136, 136), filename=str(image_paths[0]))
    first_page.insert_image(fitz.Rect(156, 36, 256, 136), filename=str(image_paths[1]))
    second_page = document.new_page()
    second_page.insert_image(fitz.Rect(36, 36, 136, 136), filename=str(image_paths[2]))
    document.save(path)
    document.close()
    return path


def _create_pdf_with_shared_image_xref(path: Path, image_path: Path) -> Path:
    document = fitz.open()
    page = document.new_page()
    xref = page.insert_image(fitz.Rect(36, 36, 136, 136), filename=str(image_path))
    page.insert_image(fitz.Rect(156, 36, 256, 136), xref=xref)
    document.save(path)
    document.close()
    return path


def test_extract_images_from_docx_uses_deterministic_names_and_order(tmp_path: Path) -> None:
    red = _create_color_image(tmp_path / "inputs" / "red.png", (255, 0, 0))
    blue = _create_color_image(tmp_path / "inputs" / "blue.png", (0, 0, 255))
    source = _create_docx_with_images(tmp_path / "sample.docx", [red, blue])
    output_dir = tmp_path / "extracted" / "docx"

    extracted = extract_images(source, output_dir)

    assert output_dir.is_dir()
    assert [path.name for path in extracted] == [
        "sample-image-0000.png",
        "sample-image-0001.png",
    ]
    assert [_image_color(path) for path in extracted] == [(255, 0, 0), (0, 0, 255)]



def test_replace_image_updates_docx_media_slot_in_place(tmp_path: Path) -> None:
    red = _create_color_image(tmp_path / "inputs" / "red.png", (255, 0, 0))
    blue = _create_color_image(tmp_path / "inputs" / "blue.png", (0, 0, 255))
    green = _create_color_image(tmp_path / "inputs" / "green.png", (0, 255, 0))
    source = _create_docx_with_images(tmp_path / "sample.docx", [red, blue])

    replace_image(source, 0, green)
    extracted = extract_images(source, tmp_path / "extracted" / "docx-replaced")

    assert [_image_color(path) for path in extracted] == [(0, 255, 0), (0, 0, 255)]



def test_extract_and_replace_pdf_images_follow_page_then_image_order(tmp_path: Path) -> None:
    red = _create_color_image(tmp_path / "inputs" / "red.png", (255, 0, 0))
    blue = _create_color_image(tmp_path / "inputs" / "blue.png", (0, 0, 255))
    green = _create_color_image(tmp_path / "inputs" / "green.png", (0, 255, 0))
    yellow = _create_color_image(tmp_path / "inputs" / "yellow.png", (255, 255, 0))
    source = _create_pdf_with_images(tmp_path / "sample.pdf", [red, blue, green])

    extracted_before = extract_images(source, tmp_path / "extracted" / "pdf-before")
    assert [path.name for path in extracted_before] == [
        "sample-image-0000.png",
        "sample-image-0001.png",
        "sample-image-0002.png",
    ]
    assert [_image_color(path) for path in extracted_before] == [
        (255, 0, 0),
        (0, 0, 255),
        (0, 255, 0),
    ]

    replace_image(source, 1, yellow)
    extracted_after = extract_images(source, tmp_path / "extracted" / "pdf-after")

    assert [_image_color(path) for path in extracted_after] == [
        (255, 0, 0),
        (255, 255, 0),
        (0, 255, 0),
    ]



def test_replace_image_rejects_pdf_slots_that_share_an_image_object(tmp_path: Path) -> None:
    red = _create_color_image(tmp_path / "inputs" / "red.png", (255, 0, 0))
    green = _create_color_image(tmp_path / "inputs" / "green.png", (0, 255, 0))
    source = _create_pdf_with_shared_image_xref(tmp_path / "shared.pdf", red)

    with pytest.raises(NotImplementedError, match=r"uniquely addressed raster image objects"):
        replace_image(source, 0, green)


@pytest.mark.parametrize("callable_obj", [extract_images, replace_image])
def test_image_helpers_validate_paths_and_replacements(tmp_path: Path, callable_obj) -> None:
    unsupported = tmp_path / "legacy.xls"
    unsupported.write_text("legacy")
    bad_replacement = tmp_path / "replacement.txt"
    bad_replacement.write_text("not an image")
    docx_source = _create_docx_with_images(
        tmp_path / "sample.docx",
        [_create_color_image(tmp_path / "inputs" / "red.png", (255, 0, 0))],
    )

    if callable_obj is extract_images:
        with pytest.raises(FileNotFoundError, match=r"Image source file '.*missing\.docx' does not exist\."):
            callable_obj(tmp_path / "missing.docx", tmp_path / "out")

        with pytest.raises(
            ValueError,
            match=r"Image source file '.*legacy\.xls' has unsupported extension 'xls'\. Supported extensions: docx, pdf, pptx, xlsm, xlsx\.",
        ):
            callable_obj(unsupported, tmp_path / "out")
    else:
        with pytest.raises(ValueError, match=r"image_index must be a zero-based non-negative integer\."):
            callable_obj(docx_source, -1, bad_replacement)

        with pytest.raises(IndexError, match=r"Image index 5 is out of range"):
            callable_obj(docx_source, 5, bad_replacement)

        with pytest.raises(ValueError, match=r"Replacement image '.*replacement\.txt' is not a readable image file\."):
            callable_obj(docx_source, 0, bad_replacement)

        with pytest.raises(
            ValueError,
            match=r"Image source file '.*legacy\.xls' has unsupported extension 'xls'\. Supported extensions: docx, pdf, pptx, xlsm, xlsx\.",
        ):
            callable_obj(unsupported, 0, bad_replacement)
