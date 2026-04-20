from __future__ import annotations

from pathlib import Path

from docx import Document
import fitz
from openpyxl import Workbook
from openpyxl.comments import Comment

from office_automation.anonymize.candidate_summary import (
    build_body_text_candidate_summary,
    resolve_body_text_confirmation,
)
from office_automation.anonymize.detect import detect
from office_automation.anonymize.transform import transform
from office_automation.anonymize.validate import validate


STRUCTURAL_CATEGORIES = {"comments", "notes", "headers", "footers"}



def _policy(
    *,
    comments: dict | None = None,
    notes: dict | None = None,
    headers: dict | None = None,
    footers: dict | None = None,
    metadata: dict | None = None,
    images: dict | None = None,
) -> dict:
    return {
        "comments": comments or {"enabled": True, "action": "remove"},
        "notes": notes or {"enabled": True, "action": "remove"},
        "headers": headers or {"enabled": True, "action": "remove"},
        "footers": footers or {"enabled": True, "action": "remove"},
        "metadata": metadata or {"enabled": True, "action": "clear"},
        "images": images or {"enabled": True, "mode": "report_only", "replacement_path": None},
        "manual_review_required": True,
    }



def _structural_findings_for(folder: Path, file_name: str) -> list[dict]:
    return [
        finding
        for finding in detect(folder)
        if finding["relative_path"] == file_name and finding["category"] in STRUCTURAL_CATEGORIES
    ]



def _findings_for(folder: Path, file_name: str, *, categories: set[str]) -> list[dict]:
    return [
        finding
        for finding in detect(folder)
        if finding["relative_path"] == file_name and finding["category"] in categories
    ]



def _findings_for_with_body_text(
    folder: Path,
    file_name: str,
    *,
    candidate_inputs: dict,
    categories: set[str] | None = None,
) -> list[dict]:
    return [
        finding
        for finding in detect(folder, body_text_candidate_inputs=candidate_inputs)
        if finding["relative_path"] == file_name and (categories is None or finding["category"] in categories)
    ]



def _create_xlsx_with_note_header_footer(path: Path) -> Path:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "SheetA"
    worksheet["A1"] = "value"
    worksheet["A1"].comment = Comment("Secret note", "Bob")
    worksheet.oddHeader.left.text = "Left Header"
    worksheet.oddFooter.center.text = "Center Footer"
    workbook.save(path)
    workbook.close()
    return path



def _create_xlsx_with_note(path: Path) -> Path:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "SheetA"
    worksheet["A1"] = "value"
    worksheet["A1"].comment = Comment("Secret note", "Bob")
    workbook.save(path)
    workbook.close()
    return path



def _create_xlsx_with_multiple_body_text_candidates(path: Path) -> Path:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Sheet1"
    worksheet["A1"] = "Primary contact: Jane Example"
    worksheet["A2"] = "Vendor: Example Corp"
    worksheet["A3"] = "Backup contact: Bob Example"
    workbook.save(path)
    workbook.close()
    return path



def _create_docx_with_comment_header_footer(path: Path) -> Path:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Document Header"
    section.footer.paragraphs[0].text = "Document Footer"
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Hello document")
    document.add_comment(run, text="Remove this comment", author="Alice", initials="AL")
    document.save(path)
    return path



def _create_pdf_with_body_text(path: Path) -> Path:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Contact john@example.com for the update")
    document.save(path)
    document.close()
    return path



def _attach_body_text_context(
    transform_results: list[dict],
    *,
    body_text_findings: list[dict],
    mode: str,
    approved_texts: list[str] | None = None,
    rejected_texts: list[str] | None = None,
    replacement_overrides: dict[str, str] | None = None,
) -> dict[str, dict]:
    summary = build_body_text_candidate_summary(body_text_findings)
    candidates_by_text = {candidate["display_text"]: candidate for candidate in summary["candidates"]}
    approved_candidate_ids = [candidates_by_text[text]["candidate_id"] for text in approved_texts or []]
    rejected_candidate_ids = [candidates_by_text[text]["candidate_id"] for text in rejected_texts or []]
    resolved_overrides = {
        candidates_by_text[text]["candidate_id"]: replacement
        for text, replacement in (replacement_overrides or {}).items()
    }
    resolution = resolve_body_text_confirmation(
        summary,
        {
            "mode": mode,
            "approved_candidate_ids": approved_candidate_ids,
            "rejected_candidate_ids": rejected_candidate_ids,
            "replacement_overrides": resolved_overrides,
            "review_notes": [],
        },
    )
    grouped_findings: dict[str, list[dict]] = {}
    for finding in body_text_findings:
        grouped_findings.setdefault(finding["relative_path"], []).append(finding)

    result_by_path = {result["relative_path"]: result for result in transform_results}
    for relative_path, findings in grouped_findings.items():
        file_summary = build_body_text_candidate_summary(findings)
        candidate_ids = {candidate["candidate_id"] for candidate in file_summary["candidates"]}
        body_text_context = {
            "candidate_summary": file_summary,
            "candidate_decisions": {
                candidate_id: payload
                for candidate_id, payload in resolution["candidate_decisions"].items()
                if candidate_id in candidate_ids
            },
            "confirmation_required": bool(file_summary["confirmation_required"]),
            "next_step_guidance": (
                "Review the body-text candidate summary and rerun with body_text_confirmation.mode=apply_confirmed."
                if mode == "preview_only"
                else "Confirmed body-text decisions were applied only to the approved subset."
            ),
            "pending_candidate_ids": [candidate_id for candidate_id in resolution["undecided_candidate_ids"] if candidate_id in candidate_ids],
            "approved_candidate_ids": [candidate_id for candidate_id in resolution["approved_candidate_ids"] if candidate_id in candidate_ids],
            "rejected_candidate_ids": [candidate_id for candidate_id in resolution["rejected_candidate_ids"] if candidate_id in candidate_ids],
            "run_mode": mode,
        }
        result_by_path[relative_path]["body_text"] = body_text_context
    return candidates_by_text



def test_validate_returns_clean_success_and_writes_markdown_report(tmp_path: Path) -> None:
    source = _create_xlsx_with_note(tmp_path / "cleanable.xlsx")
    findings = _findings_for(tmp_path, source.name, categories={"notes", "metadata"})
    transform_results = transform(
        findings,
        _policy(
            notes={"enabled": True, "action": "clear"},
            metadata={"enabled": True, "action": "clear"},
        ),
    )

    results = validate(str(tmp_path), transform_results)

    assert len(results) == 1
    result = results[0]
    report_path = tmp_path / "anonymization_report.md"

    assert result["file_path"] == str(source)
    assert result["report_path"] == str(report_path)
    assert result["status"] == "success"
    assert result["residual_findings"] == []
    assert result["warnings"] == []
    assert result["manual_review_items"] == []
    assert result["body_text"] == {
        "approved_candidate_ids": [],
        "candidate_summary_count": 0,
        "confirmation_required": False,
        "decision_counts": {
            "approved": 0,
            "low_confidence": 0,
            "manual_review_required": 0,
            "pending": 0,
            "rejected": 0,
            "residual": 0,
            "skipped": 0,
        },
        "decision_trace": [],
        "low_confidence_candidate_ids": [],
        "next_step_guidance": "No body-text candidates were generated for this run.",
        "pending_candidate_ids": [],
        "rejected_candidate_ids": [],
        "residual_candidate_ids": [],
        "run_mode": "absent",
    }
    assert result["transform_summary"]["transform_status"] == "success"
    assert result["transform_summary"]["action_status_counts"] == {"applied": 2}
    assert {action["validation_outcome"] for action in result["transform_summary"]["actions"]} == {"cleared"}

    assert report_path.exists()
    report = report_path.read_text(encoding="utf-8")
    assert "# Anonymization Report" in report
    assert "## Run summary" in report
    assert "## Body-text candidate summary" in report
    assert "- No body-text candidates were generated for this run." in report
    assert "## Body-text decision trace and next-step guidance" in report
    assert "## Residual findings after re-scan" in report
    assert "`cleanable.xlsx`" in report
    assert "Overall run outcome: `success`" in report



def test_validate_preserves_residual_and_manual_review_context_in_report(tmp_path: Path) -> None:
    source = _create_docx_with_comment_header_footer(tmp_path / "commented.docx")
    findings = _findings_for(tmp_path, source.name, categories=STRUCTURAL_CATEGORIES | {"metadata"})
    transform_results = transform(
        findings,
        _policy(
            comments={"enabled": True, "action": "remove"},
            headers={"enabled": True, "action": "replace", "replacement_text": "Sanitized Header"},
            footers={"enabled": True, "action": "clear"},
            metadata={"enabled": True, "action": "clear"},
        ),
    )

    results = validate(str(tmp_path), transform_results)

    assert len(results) == 1
    result = results[0]
    report_path = tmp_path / "anonymization_report.md"

    assert result["status"] == "partial_success"
    assert result["report_path"] == str(report_path)
    assert [residual["category"] for residual in result["residual_findings"]] == ["comments"]
    assert result["residual_findings"][0]["validation_outcome"] == "manual_review_required"
    assert "requiring manual review" in result["residual_findings"][0]["reason"].lower()

    action_outcomes = {action["category"]: action["validation_outcome"] for action in result["transform_summary"]["actions"]}
    assert action_outcomes == {
        "comments": "manual_review_required",
        "footers": "cleared",
        "headers": "changed_in_place",
        "metadata": "cleared",
    }

    manual_review_categories = {item["category"] for item in result["manual_review_items"]}
    assert {"comments", "headers"}.issubset(manual_review_categories)
    assert result["warnings"] == []

    assert report_path.exists()
    report = report_path.read_text(encoding="utf-8")
    assert "## Manual-review checklist" in report
    assert "## Residual findings after re-scan" in report
    assert "`commented.docx`" in report
    assert "category `comments`" in report
    assert "Overall run outcome: `partial_success`" in report



def test_validate_marks_preview_only_body_text_runs_as_manual_review_and_writes_replay_guidance(tmp_path: Path) -> None:
    source = _create_xlsx_with_multiple_body_text_candidates(tmp_path / "preview.xlsx")
    body_text_findings = _findings_for_with_body_text(
        tmp_path,
        source.name,
        candidate_inputs={"person_names": ["Jane Example", "Bob Example"], "company_names": ["Example Corp"]},
        categories={"body_text"},
    )
    metadata_findings = _findings_for(tmp_path, source.name, categories={"metadata"})
    transform_results = transform(
        metadata_findings,
        _policy(metadata={"enabled": True, "action": "clear"}),
    )
    candidates_by_text = _attach_body_text_context(
        transform_results,
        body_text_findings=body_text_findings,
        mode="preview_only",
    )

    results = validate(str(tmp_path), transform_results)

    assert len(results) == 1
    result = results[0]
    report_path = tmp_path / "anonymization_report.md"

    assert result["status"] == "manual_review_required"
    assert result["body_text"]["run_mode"] == "preview_only"
    assert result["body_text"]["confirmation_required"] is True
    assert result["body_text"]["candidate_summary_count"] == 3
    assert result["body_text"]["approved_candidate_ids"] == []
    assert result["body_text"]["rejected_candidate_ids"] == []
    assert result["body_text"]["pending_candidate_ids"] == [
        candidates_by_text["Bob Example"]["candidate_id"],
        candidates_by_text["Example Corp"]["candidate_id"],
        candidates_by_text["Jane Example"]["candidate_id"],
    ]
    assert result["body_text"]["decision_counts"] == {
        "approved": 0,
        "low_confidence": 0,
        "manual_review_required": 0,
        "pending": 3,
        "rejected": 0,
        "residual": 0,
        "skipped": 0,
    }
    assert {entry["decision"] for entry in result["body_text"]["decision_trace"]} == {"pending"}
    assert result["body_text"]["next_step_guidance"].endswith("body_text_confirmation.mode=apply_confirmed.")

    assert report_path.exists()
    report = report_path.read_text(encoding="utf-8")
    assert "## Body-text candidate summary" in report
    assert "## Body-text decision trace and next-step guidance" in report
    assert "body_text_run_mode: `preview_only`" in report
    assert "body_text_confirmation.mode=apply_confirmed" in report
    assert candidates_by_text["Jane Example"]["candidate_id"] in report



def test_validate_run_summary_deduplicates_body_text_candidate_ids_across_files(tmp_path: Path) -> None:
    source_a = _create_xlsx_with_multiple_body_text_candidates(tmp_path / "preview-a.xlsx")
    source_b = _create_xlsx_with_multiple_body_text_candidates(tmp_path / "preview-b.xlsx")
    body_text_findings = _findings_for_with_body_text(
        tmp_path,
        source_a.name,
        candidate_inputs={"person_names": ["Jane Example", "Bob Example"], "company_names": ["Example Corp"]},
        categories={"body_text"},
    ) + _findings_for_with_body_text(
        tmp_path,
        source_b.name,
        candidate_inputs={"person_names": ["Jane Example", "Bob Example"], "company_names": ["Example Corp"]},
        categories={"body_text"},
    )
    metadata_findings = _findings_for(tmp_path, source_a.name, categories={"metadata"}) + _findings_for(
        tmp_path,
        source_b.name,
        categories={"metadata"},
    )
    transform_results = transform(
        metadata_findings,
        _policy(metadata={"enabled": True, "action": "clear"}),
    )
    _attach_body_text_context(
        transform_results,
        body_text_findings=body_text_findings,
        mode="preview_only",
    )

    results = validate(str(tmp_path), transform_results)

    assert len(results) == 2
    assert sum(result["body_text"]["candidate_summary_count"] for result in results) == 6

    report = (tmp_path / "anonymization_report.md").read_text(encoding="utf-8")

    assert "- body_text_candidate_summary_count: 3" in report
    assert (
        '- body_text_decision_counts: `{"approved": 0, "low_confidence": 0, '
        '"manual_review_required": 0, "pending": 3, "rejected": 0, "residual": 0, "skipped": 0}`'
    ) in report



def test_validate_reports_confirmed_body_text_cleared_residual_and_rejected_candidates(tmp_path: Path) -> None:
    source = _create_xlsx_with_multiple_body_text_candidates(tmp_path / "confirmed.xlsx")
    findings = _findings_for_with_body_text(
        tmp_path,
        source.name,
        candidate_inputs={"person_names": ["Jane Example", "Bob Example"], "company_names": ["Example Corp"]},
        categories={"body_text"},
    )
    body_text_findings = [finding for finding in findings if finding["category"] == "body_text"]
    summary = build_body_text_candidate_summary(body_text_findings)
    candidates_by_text = {candidate["display_text"]: candidate for candidate in summary["candidates"]}
    resolution = resolve_body_text_confirmation(
        summary,
        {
            "mode": "apply_confirmed",
            "approved_candidate_ids": [
                candidates_by_text["Bob Example"]["candidate_id"],
                candidates_by_text["Jane Example"]["candidate_id"],
            ],
            "rejected_candidate_ids": [candidates_by_text["Example Corp"]["candidate_id"]],
            "replacement_overrides": {
                candidates_by_text["Bob Example"]["candidate_id"]: "[PERSON]",
                candidates_by_text["Jane Example"]["candidate_id"]: "Jane Example",
            },
            "review_notes": [],
        },
    )
    transform_results = transform(
        findings,
        {
            **_policy(),
            "body_text": {
                "enabled": True,
                "mode": "apply_confirmed",
                "approved_candidate_ids": list(resolution["approved_candidate_ids"]),
                "rejected_candidate_ids": list(resolution["rejected_candidate_ids"]),
                "undecided_candidate_ids": list(resolution["undecided_candidate_ids"]),
                "approved_finding_ids": list(resolution["approved_finding_ids"]),
                "non_transformable_finding_ids": list(resolution["non_transformable_finding_ids"]),
                "candidate_decisions": dict(resolution["candidate_decisions"]),
                "candidate_summary": {
                    "candidates": list(summary["candidates"]),
                    "finding_to_candidate": dict(summary["finding_to_candidate"]),
                },
                "replacement_text": "[BODY TEXT REDACTED]",
                "replacement_map": {},
                "replacement_overrides": {
                    candidates_by_text["Bob Example"]["candidate_id"]: "[PERSON]",
                    candidates_by_text["Jane Example"]["candidate_id"]: "Jane Example",
                },
                "confirmation_warnings": list(resolution["warnings"]),
            },
        },
    )
    _attach_body_text_context(
        transform_results,
        body_text_findings=body_text_findings,
        mode="apply_confirmed",
        approved_texts=["Bob Example", "Jane Example"],
        rejected_texts=["Example Corp"],
        replacement_overrides={"Bob Example": "[PERSON]", "Jane Example": "Jane Example"},
    )

    results = validate(str(tmp_path), transform_results)

    assert len(results) == 1
    result = results[0]
    report = (tmp_path / "anonymization_report.md").read_text(encoding="utf-8")

    assert result["status"] == "partial_success"
    assert result["body_text"]["run_mode"] == "apply_confirmed"
    assert result["body_text"]["approved_candidate_ids"] == [
        candidates_by_text["Bob Example"]["candidate_id"],
        candidates_by_text["Jane Example"]["candidate_id"],
    ]
    assert result["body_text"]["rejected_candidate_ids"] == [candidates_by_text["Example Corp"]["candidate_id"]]
    assert result["body_text"]["pending_candidate_ids"] == []
    assert result["body_text"]["residual_candidate_ids"] == [candidates_by_text["Jane Example"]["candidate_id"]]
    assert result["body_text"]["low_confidence_candidate_ids"] == []
    assert result["body_text"]["decision_counts"] == {
        "approved": 2,
        "low_confidence": 0,
        "manual_review_required": 0,
        "pending": 0,
        "rejected": 1,
        "residual": 1,
        "skipped": 1,
    }
    trace_by_candidate = {entry["candidate_id"]: entry for entry in result["body_text"]["decision_trace"]}
    assert trace_by_candidate[candidates_by_text["Bob Example"]["candidate_id"]]["validation_outcome"] == "cleared"
    assert trace_by_candidate[candidates_by_text["Jane Example"]["candidate_id"]]["validation_outcome"] == "residual"
    assert trace_by_candidate[candidates_by_text["Example Corp"]["candidate_id"]]["validation_outcome"] == "rejected_skip"
    assert candidates_by_text["Jane Example"]["candidate_id"] in report
    assert candidates_by_text["Example Corp"]["candidate_id"] in report
    assert "residual candidate ids" in report.lower()
    assert "rejected_skip" in report



def test_validate_surfaces_low_confidence_pdf_body_text_candidates_for_manual_review(tmp_path: Path) -> None:
    source = _create_pdf_with_body_text(tmp_path / "body.pdf")
    findings = _findings_for_with_body_text(
        tmp_path,
        source.name,
        candidate_inputs={"emails": ["john@example.com"]},
        categories={"body_text"},
    )
    summary = build_body_text_candidate_summary(findings)
    candidate = summary["candidates"][0]
    resolution = resolve_body_text_confirmation(
        summary,
        {
            "mode": "apply_confirmed",
            "approved_candidate_ids": [candidate["candidate_id"]],
            "rejected_candidate_ids": [],
            "replacement_overrides": {candidate["candidate_id"]: "[EMAIL]"},
            "review_notes": [],
        },
    )
    transform_results = transform(
        findings,
        {
            **_policy(),
            "body_text": {
                "enabled": True,
                "mode": "apply_confirmed",
                "approved_candidate_ids": list(resolution["approved_candidate_ids"]),
                "rejected_candidate_ids": list(resolution["rejected_candidate_ids"]),
                "undecided_candidate_ids": list(resolution["undecided_candidate_ids"]),
                "approved_finding_ids": list(resolution["approved_finding_ids"]),
                "non_transformable_finding_ids": list(resolution["non_transformable_finding_ids"]),
                "candidate_decisions": dict(resolution["candidate_decisions"]),
                "candidate_summary": {
                    "candidates": list(summary["candidates"]),
                    "finding_to_candidate": dict(summary["finding_to_candidate"]),
                },
                "replacement_text": "[BODY TEXT REDACTED]",
                "replacement_map": {},
                "replacement_overrides": {candidate["candidate_id"]: "[EMAIL]"},
                "confirmation_warnings": list(resolution["warnings"]),
            },
        },
    )
    _attach_body_text_context(
        transform_results,
        body_text_findings=findings,
        mode="apply_confirmed",
        approved_texts=["john@example.com"],
        replacement_overrides={"john@example.com": "[EMAIL]"},
    )

    results = validate(str(tmp_path), transform_results)

    assert len(results) == 1
    result = results[0]
    report = (tmp_path / "anonymization_report.md").read_text(encoding="utf-8")

    assert result["status"] == "manual_review_required"
    assert result["body_text"]["run_mode"] == "apply_confirmed"
    assert result["body_text"]["approved_candidate_ids"] == [candidate["candidate_id"]]
    assert result["body_text"]["low_confidence_candidate_ids"] == [candidate["candidate_id"]]
    assert result["body_text"]["decision_counts"] == {
        "approved": 1,
        "low_confidence": 1,
        "manual_review_required": 1,
        "pending": 0,
        "rejected": 0,
        "residual": 0,
        "skipped": 0,
    }
    assert result["body_text"]["decision_trace"][0]["validation_outcome"] == "low_confidence_manual_review"
    assert any(item["category"] == "body_text" for item in result["manual_review_items"])
    assert "pdf text-layer matches remain review-first" in report.lower()
    assert "## Manual-review checklist" in report
