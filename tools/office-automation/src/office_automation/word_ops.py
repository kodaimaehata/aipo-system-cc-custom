"""Word V1 read/edit helpers centered on body paragraphs and tables.

Supported edit operations are applied sequentially against the current document state:
- replace_paragraph_text
- insert_paragraph_after
- delete_paragraph
- replace_table_cell
- append_table_row
- update_table_row

This module is intentionally limited to `.docx` files and ordinary document-body
paragraph/table content. Track changes cleanup, text-box-driven editing, and other
full-fidelity Word surgery remain outside SG2 V1 scope.
"""

from __future__ import annotations

from collections.abc import Sequence
from os import PathLike
from pathlib import Path
import shutil
import warnings
import xml.etree.ElementTree as ET
import zipfile

from office_automation.common.files import copy_original

_SUPPORTED_EXTENSION = ".docx"
_UNSUPPORTED_OPERATION_MESSAGES = {
    "accept_track_changes": "Track changes acceptance/rejection is outside Word V1 runtime scope.",
    "reject_track_changes": "Track changes acceptance/rejection is outside Word V1 runtime scope.",
    "cleanup_track_changes": "Track changes acceptance/rejection is outside Word V1 runtime scope.",
    "edit_text_box": "Complex text-box editing is outside Word V1 runtime scope.",
    "replace_text_box_text": "Complex text-box editing is outside Word V1 runtime scope.",
}
_UNSUPPORTED_OPTION_MESSAGES = {
    "accept_track_changes": "Track changes acceptance/rejection is outside Word V1 runtime scope.",
    "reject_track_changes": "Track changes acceptance/rejection is outside Word V1 runtime scope.",
    "cleanup_track_changes": "Track changes acceptance/rejection is outside Word V1 runtime scope.",
    "edit_text_boxes": "Complex text-box editing is outside Word V1 runtime scope.",
}
_TRACK_CHANGE_LOCAL_NAMES = {
    "ins",
    "del",
    "moveFrom",
    "moveFromRangeEnd",
    "moveFromRangeStart",
    "moveTo",
    "moveToRangeEnd",
    "moveToRangeStart",
}


def read(file_path: str | PathLike[str] | Path) -> dict:
    """Read a `.docx` file and return ordered body paragraph/table data."""
    source = _validate_source_path(file_path)
    load_document = _load_document_callable()
    document = load_document(source)
    feature_warnings = _scan_document_feature_warnings(source)

    paragraphs: list[dict] = []
    tables: list[dict] = []
    body: list[dict] = []

    for body_index, block in enumerate(_iter_document_body(document)):
        if _is_paragraph(block):
            paragraph_index = len(paragraphs)
            paragraph_entry = {
                "paragraph_index": paragraph_index,
                "body_index": body_index,
                "text": block.text,
                "style": _get_style_name(block),
            }
            paragraphs.append(paragraph_entry)
            body.append(
                {
                    "body_index": body_index,
                    "type": "paragraph",
                    "paragraph_index": paragraph_index,
                    "text": block.text,
                }
            )
            continue

        table_index = len(tables)
        row_entries: list[dict] = []
        max_column_count = 0
        for row_index, row in enumerate(block.rows):
            cell_entries: list[dict] = []
            for column_index, cell in enumerate(row.cells):
                cell_entries.append(
                    {
                        "row_index": row_index,
                        "column_index": column_index,
                        "text": _cell_text(cell),
                    }
                )
            row_entries.append({"row_index": row_index, "cells": cell_entries})
            max_column_count = max(max_column_count, len(cell_entries))

        table_entry = {
            "table_index": table_index,
            "body_index": body_index,
            "row_count": len(row_entries),
            "column_count": max_column_count,
            "rows": row_entries,
        }
        tables.append(table_entry)
        body.append(
            {
                "body_index": body_index,
                "type": "table",
                "table_index": table_index,
                "row_count": len(row_entries),
                "column_count": max_column_count,
            }
        )

    return {
        "format": "docx",
        "file_path": str(source),
        "metadata": _read_core_properties(document),
        "paragraphs": paragraphs,
        "tables": tables,
        "body": body,
        "warnings": feature_warnings,
    }


def edit(file_path: str | PathLike[str] | Path, instructions: dict) -> Path:
    """Apply supported Word V1 edits and return the saved `.docx` path."""
    source = _validate_source_path(file_path)
    instruction_dict = _validate_instruction_envelope(instructions)
    operations = instruction_dict["operations"]
    options = instruction_dict.get("options", {})

    _reject_unsupported_requests(operations=operations, options=options)

    feature_warnings = _scan_document_feature_warnings(source)
    for message in feature_warnings:
        warnings.warn(message, RuntimeWarning, stacklevel=2)

    target_path, use_copy_helper = _resolve_edit_target(source, instruction_dict)
    load_document = _load_document_callable()

    if use_copy_helper:
        staged_copy = copy_original(source, target_path.parent)
        if staged_copy != target_path:
            shutil.move(str(staged_copy), str(target_path))
        document = load_document(target_path)
        save_path = target_path
    elif _paths_refer_to_same_location(source, target_path):
        document = load_document(source)
        save_path = source
    elif instruction_dict["copy_before_edit"]:
        shutil.copy2(source, target_path)
        document = load_document(target_path)
        save_path = target_path
    else:
        document = load_document(source)
        save_path = target_path

    emitted_format_warning = False
    for operation in operations:
        emitted_format_warning = _apply_operation(
            document,
            operation,
            emitted_format_warning=emitted_format_warning,
        )

    document.save(save_path)
    return Path(save_path)


def _load_document_callable():
    try:
        from docx import Document as load_document
    except ModuleNotFoundError as exc:
        raise ModuleNotFoundError(
            "python-docx is required for Word operations. Install the office-edit dependencies."
        ) from exc
    return load_document


def _load_word_types():
    try:
        from docx.document import Document as DocumentType
        from docx.oxml import OxmlElement
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ModuleNotFoundError as exc:
        raise ModuleNotFoundError(
            "python-docx is required for Word operations. Install the office-edit dependencies."
        ) from exc
    return DocumentType, OxmlElement, CT_P, CT_Tbl, Paragraph, Table


def _validate_source_path(file_path: str | PathLike[str] | Path) -> Path:
    source = Path(file_path)
    if not source.exists():
        raise FileNotFoundError(f"Word source file '{source}' does not exist.")
    if source.is_symlink() or not source.is_file():
        raise ValueError(f"Word source file '{source}' is not a regular file.")
    if source.suffix.lower() != _SUPPORTED_EXTENSION:
        display_extension = source.suffix.lower().lstrip(".") or "<none>"
        raise ValueError(
            f"Word source file '{source}' has unsupported extension '{display_extension}'. "
            "Supported extension: docx."
        )
    return source


def _validate_instruction_envelope(instructions: dict) -> dict:
    if not isinstance(instructions, dict):
        raise TypeError("Word edit instructions must be a dict.")

    operations = instructions.get("operations")
    if not isinstance(operations, list) or not operations:
        raise ValueError("Word edit instructions must include a non-empty 'operations' list.")
    for operation in operations:
        if not isinstance(operation, dict):
            raise TypeError("Each Word edit operation must be a dict.")

    output_path = instructions.get("output_path")
    if output_path in (None, ""):
        raise ValueError(
            "Word edit instructions must include 'output_path'; runtime save-path fallback is not supported."
        )

    options = instructions.get("options", {})
    if not isinstance(options, dict):
        raise TypeError("Word edit instructions 'options' must be a dict when provided.")

    copy_before_edit = instructions.get("copy_before_edit", True)
    if not isinstance(copy_before_edit, bool):
        raise TypeError("Word edit instructions 'copy_before_edit' must be a bool.")

    envelope = dict(instructions)
    envelope["operations"] = operations
    envelope["output_path"] = output_path
    envelope["options"] = options
    envelope["copy_before_edit"] = copy_before_edit
    return envelope


def _reject_unsupported_requests(*, operations: list[dict], options: dict) -> None:
    for operation in operations:
        operation_name = operation.get("operation")
        if operation_name in _UNSUPPORTED_OPERATION_MESSAGES:
            raise NotImplementedError(_UNSUPPORTED_OPERATION_MESSAGES[operation_name])

    for option_name, message in _UNSUPPORTED_OPTION_MESSAGES.items():
        if options.get(option_name):
            raise NotImplementedError(message)


def _resolve_edit_target(source: Path, instructions: dict) -> tuple[Path, bool]:
    output_path = instructions["output_path"]
    copy_before_edit = instructions["copy_before_edit"]

    target_path = Path(output_path)
    _validate_target_extension(target_path)
    _prepare_output_path(source, target_path, copy_before_edit=copy_before_edit)

    use_copy_helper = (
        copy_before_edit
        and target_path.name == source.name
        and target_path.parent != source.parent
    )
    return target_path, use_copy_helper


def _prepare_output_path(source: Path, target_path: Path, *, copy_before_edit: bool) -> None:
    parent = target_path.parent
    if parent.exists() and not parent.is_dir():
        raise NotADirectoryError(f"Word output directory '{parent}' is not a directory.")
    parent.mkdir(parents=True, exist_ok=True)

    if copy_before_edit and _paths_refer_to_same_location(source, target_path):
        raise ValueError(
            "Word edit instructions cannot preserve the original when 'output_path' "
            "points to the source file and 'copy_before_edit' is True."
        )

    if target_path.exists() and not _paths_refer_to_same_location(source, target_path):
        raise FileExistsError(
            f"Word output path '{target_path}' already exists; the runtime will not silently overwrite it."
        )


def _validate_target_extension(path: Path) -> None:
    if path.suffix.lower() != _SUPPORTED_EXTENSION:
        display_extension = path.suffix.lower().lstrip(".") or "<none>"
        raise ValueError(
            f"Word output path '{path}' has unsupported extension '{display_extension}'. "
            "Supported extension: docx."
        )


def _apply_operation(document, operation: dict, *, emitted_format_warning: bool) -> bool:
    operation_name = operation.get("operation")
    if not isinstance(operation_name, str) or not operation_name:
        raise ValueError("Word edit operations must include a non-empty 'operation' string.")

    if operation_name == "replace_paragraph_text":
        paragraph = _get_paragraph(document, _require_non_negative_int(operation, "paragraph_index"))
        paragraph.text = _require_text(operation, "text")
        return _emit_format_warning_once(emitted_format_warning)

    if operation_name == "insert_paragraph_after":
        paragraph = _get_paragraph(document, _require_non_negative_int(operation, "paragraph_index"))
        new_paragraph = _insert_paragraph_after(
            paragraph,
            text=_require_text(operation, "text"),
            style_name=operation.get("style"),
        )
        if style_name := operation.get("style"):
            new_paragraph.style = style_name
        return _emit_format_warning_once(emitted_format_warning)

    if operation_name == "delete_paragraph":
        paragraph = _get_paragraph(document, _require_non_negative_int(operation, "paragraph_index"))
        _delete_paragraph(paragraph)
        return emitted_format_warning

    if operation_name == "replace_table_cell":
        table = _get_table(document, _require_non_negative_int(operation, "table_index"))
        row_index = _require_non_negative_int(operation, "row_index")
        column_index = _require_non_negative_int(operation, "column_index")
        cell = _get_table_cell(table, row_index=row_index, column_index=column_index)
        cell.text = _require_text(operation, "text")
        return _emit_format_warning_once(emitted_format_warning)

    if operation_name == "append_table_row":
        table = _get_table(document, _require_non_negative_int(operation, "table_index"))
        values = _require_row_values(operation)
        row = table.add_row()
        _write_cells(row.cells, values)
        return _emit_format_warning_once(emitted_format_warning)

    if operation_name == "update_table_row":
        table = _get_table(document, _require_non_negative_int(operation, "table_index"))
        row_index = _require_non_negative_int(operation, "row_index")
        values = _require_row_values(operation)
        row = _get_table_row(table, row_index)
        _write_cells(row.cells, values)
        return _emit_format_warning_once(emitted_format_warning)

    raise ValueError(f"Unsupported Word edit operation '{operation_name}'.")


def _emit_format_warning_once(already_emitted: bool) -> bool:
    if already_emitted:
        return True
    warnings.warn(
        "Word V1 text edits operate at paragraph/cell granularity and may normalize "
        "run-level formatting.",
        RuntimeWarning,
        stacklevel=3,
    )
    return True


def _iter_document_body(document):
    DocumentType, _, CT_P, CT_Tbl, Paragraph, Table = _load_word_types()
    if not isinstance(document, DocumentType):
        raise TypeError("Expected a python-docx Document instance.")

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _is_paragraph(block) -> bool:
    return block.__class__.__name__ == "Paragraph"


def _get_style_name(paragraph) -> str | None:
    style = getattr(paragraph, "style", None)
    return getattr(style, "name", None)


def _cell_text(cell) -> str:
    texts = [paragraph.text for paragraph in cell.paragraphs]
    if not texts:
        return cell.text
    return "\n".join(texts)


def _read_core_properties(document) -> dict:
    properties = document.core_properties
    return {
        "title": properties.title,
        "author": properties.author,
        "subject": properties.subject,
        "category": properties.category,
        "comments": properties.comments,
        "keywords": properties.keywords,
        "last_modified_by": properties.last_modified_by,
        "revision": properties.revision,
        "created": _serialize_datetime(properties.created),
        "modified": _serialize_datetime(properties.modified),
    }


def _serialize_datetime(value) -> str | None:
    if value is None:
        return None
    try:
        return value.isoformat()
    except AttributeError:
        return str(value)


def _scan_document_feature_warnings(path: Path) -> list[str]:
    warning_messages: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            has_track_changes = False
            has_text_boxes = False
            for member_name in archive.namelist():
                if not member_name.startswith("word/") or not member_name.endswith(".xml"):
                    continue
                try:
                    root = ET.fromstring(archive.read(member_name))
                except ET.ParseError:
                    continue
                for element in root.iter():
                    local_name = _xml_local_name(element.tag)
                    if local_name in _TRACK_CHANGE_LOCAL_NAMES:
                        has_track_changes = True
                    elif local_name == "txbxContent":
                        has_text_boxes = True
                    if has_track_changes and has_text_boxes:
                        break
                if has_track_changes and has_text_boxes:
                    break
    except (FileNotFoundError, zipfile.BadZipFile):
        return warning_messages

    if has_track_changes:
        warning_messages.append(
            "This document contains Word revision/track-change markup. SG2 Word V1 does "
            "not accept or reject track changes."
        )
    if has_text_boxes:
        warning_messages.append(
            "This document contains Word text box content. SG2 Word V1 only supports body "
            "paragraph and table edits."
        )
    return warning_messages


def _xml_local_name(tag: str) -> str:
    return tag.split("}", 1)[-1]


def _get_paragraph(document, paragraph_index: int):
    paragraphs = list(document.paragraphs)
    try:
        return paragraphs[paragraph_index]
    except IndexError as exc:
        raise IndexError(
            f"Paragraph index {paragraph_index} is out of range for this document."
        ) from exc


def _get_table(document, table_index: int):
    tables = list(document.tables)
    try:
        return tables[table_index]
    except IndexError as exc:
        raise IndexError(f"Table index {table_index} is out of range for this document.") from exc


def _get_table_row(table, row_index: int):
    rows = list(table.rows)
    try:
        return rows[row_index]
    except IndexError as exc:
        raise IndexError(f"Table row index {row_index} is out of range for this table.") from exc


def _get_table_cell(table, *, row_index: int, column_index: int):
    row = _get_table_row(table, row_index)
    cells = list(row.cells)
    try:
        return cells[column_index]
    except IndexError as exc:
        raise IndexError(
            f"Table cell column index {column_index} is out of range for row {row_index}."
        ) from exc


def _write_cells(cells, values: Sequence[str]) -> None:
    if len(values) > len(cells):
        raise ValueError(
            f"Row operation provided {len(values)} values for a table row with {len(cells)} cells."
        )
    for index, cell in enumerate(cells):
        cell.text = values[index] if index < len(values) else ""


def _require_non_negative_int(operation: dict, key: str) -> int:
    value = operation.get(key)
    if not isinstance(value, int) or isinstance(value, bool) or value < 0:
        raise TypeError(f"Word edit operation '{key}' must be a non-negative integer.")
    return value


def _require_text(operation: dict, key: str) -> str:
    value = operation.get(key)
    if not isinstance(value, str):
        raise TypeError(f"Word edit operation '{key}' must be a string.")
    return value


def _require_row_values(operation: dict) -> list[str]:
    values = operation.get("values")
    if isinstance(values, str) or not isinstance(values, Sequence):
        raise TypeError("Word row operations must include 'values' as a sequence of strings.")

    normalized_values: list[str] = []
    for value in values:
        if not isinstance(value, str):
            raise TypeError("Word row operation values must all be strings.")
        normalized_values.append(value)
    return normalized_values


def _insert_paragraph_after(paragraph, *, text: str, style_name: str | None):
    _, OxmlElement, _, _, Paragraph, _ = _load_word_types()
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_paragraph.style = style_name
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is None:
        raise ValueError("Cannot delete a paragraph without a parent XML node.")
    parent.remove(element)


def _paths_refer_to_same_location(left: Path, right: Path) -> bool:
    return left.resolve() == right.resolve(strict=False)
